import streamlit as st
import os
import shutil
import tempfile
import base64
import uuid
import pandas as pd
from langdetect import detect
from PyPDF2 import PdfReader
from gtts import gTTS
from docx import Document as DocxDocument
from langchain_community.llms import Ollama
from langchain.chains.question_answering import load_qa_chain
from langchain.docstore.document import Document

# === Updated Custom CSS with Centered Chat and No Spacing ===
st.markdown("""
<style>
    /* Reset margins and padding */
    body, html, .stApp {
        margin: 0;
        padding: 0;
        overflow: hidden;
    }
    
    .main-container {
        display: flex;
        flex-direction: column;
    }
    
    /* Title container with minimal spacing */
    .title-container {
        margin-bottom: 0 !important;
        padding-bottom: 0 !important;
    }
    
    .title-container h1 {
        margin-bottom: 0 !important;
        padding-bottom: 0 !important;
    }
    
    .title-container .stMarkdown {
        margin-top: 0 !important;
        padding-top: 0 !important;
        margin-bottom: 5px !important;
    }
    
    /* Input container with minimal spacing */
    .input-container {
        background-color: #1e293b;
        padding: 10px;
        border-radius: 8px;
        margin-top: 5px;
        margin-bottom: 0;  /* Remove bottom margin */
        z-index: 10;
    }
    
    /* Chat history container - remove top margin */
    .chat-history-container {
        flex: 1;
        background-color: #1e293b;
        border-radius: 8px;
        display: flex;
        flex-direction: column;
        overflow: hidden;
        margin-top: 0;  /* Remove top margin */
    }

    .chat-history {
        flex: 1;
        overflow-y: auto;
        padding: 10px;
        background-color: #1e293b;
        display: flex;
        flex-direction: column;
        align-items: center;  /* Center messages horizontally */
    }

    .message {
        padding: 10px;
        margin-bottom: 10px;
        border-radius: 8px;
        width: 90%;  /* Consistent width */
        max-width: 800px;  /* Maximum width for larger screens */
    }

    .user-message {
        background-color: #2d3748;
        border-left: 4px solid #4f46e5;
    }

    .assistant-message {
        background-color: #1e3a8a;
        border-left: 4px solid #10b981;
    }

    .message-header {
        display: flex;
        margin-bottom: 5px;
        font-weight: bold;
        justify-content: center;  /* Center header */
    }

    .chat-history::-webkit-scrollbar {
        width: 6px;
    }

    .chat-history::-webkit-scrollbar-track {
        background: #1e293b;
    }

    .chat-history::-webkit-scrollbar-thumb {
        background: #4f46e5;
        border-radius: 3px;
    }

    .stButton>button {
        background-color: #4f46e5 !important;
        color: white !important;
        border-radius: 6px !important;
        padding: 0.4rem 0.8rem !important;
        width: 100%;
        font-size: 0.9em;
    }

    .stTextInput>div>div>input {
        background-color: #1e293b !important;
        color: white !important;
        border-radius: 6px !important;
        padding: 0.6rem !important;
    }
    
    /* Column layout for buttons */
    .button-columns {
        display: flex;
        gap: 10px;
        margin-top: 10px;
        margin-bottom: 0;  /* Remove bottom margin */
    }
    
    .button-columns > * {
        flex: 1;
    }
</style>
""", unsafe_allow_html=True)

# === Initialize Session State ===
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "last_answer" not in st.session_state:
    st.session_state.last_answer = ""
if "play_audio" not in st.session_state:
    st.session_state.play_audio = False
if "question" not in st.session_state:
    st.session_state.question = ""
if "file_preview" not in st.session_state:
    st.session_state.file_preview = ""
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None
if "selected_files" not in st.session_state:
    st.session_state.selected_files = []
if "folder_path" not in st.session_state:
    st.session_state.folder_path = r"YourFolderPath"

# === Helper Functions ===
def play_gtts(text, lang_code="en"):
    try:
        tts = gTTS(text=text, lang=lang_code)
        tmp_file_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.mp3")
        tts.save(tmp_file_path)
        with open(tmp_file_path, "rb") as audio_file:
            audio_data = audio_file.read()
        st.audio(audio_data, format="audio/mp3")
        b64 = base64.b64encode(audio_data).decode()
        href = f'<a href="data:audio/mp3;base64,{b64}" download="audio.mp3" class="stButton">⬇️ Download Audio</a>'
        st.markdown(href, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"❌ TTS failed: {e}")

def get_file_icon(ext):
    icons = {
        ".pdf": "📄",
        ".docx": "📝",
        ".txt": "📃",
        ".csv": "📊",
        ".xlsx": "📈"
    }
    return icons.get(ext.lower(), "📁")

def extract_text_from_file(filepath):
    try:
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".pdf":
            with open(filepath, "rb") as f:
                reader = PdfReader(f)
                return "\n".join([p.extract_text() or "" for p in reader.pages])
        elif ext == ".docx":
            doc = DocxDocument(filepath)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        elif ext in [".csv", ".xlsx"]:
            if ext == ".csv":
                df = pd.read_csv(filepath)
            else:
                df = pd.read_excel(filepath)
            return df.to_string(index=False)
    except Exception as e:
        st.warning(f"⚠️ Could not read {filepath}: {e}")
    return ""

def load_doc(filepath):
    text = extract_text_from_file(filepath)
    if text.strip():
        return Document(page_content=text, metadata={"source": os.path.basename(filepath)})
    return None

# === Page Setup ===
st.set_page_config(
    page_title="📚 AyF - Ask Your Folder",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# === Sidebar - Setup Controls ===
with st.sidebar:
    st.title("📚 Ask your Folder")
    st.caption("AyF 0.1 Beta")
    
    # Folder Setup
    with st.expander("📁 Folder Setup", expanded=True):
        folder_path = st.text_input("Document Folder Path:", 
                                   value=st.session_state.folder_path,
                                   key="folder_path_input")
        
        # Create folder if it doesn't exist
        os.makedirs(folder_path, exist_ok=True)
        st.session_state.folder_path = folder_path
        
        st.markdown("**Add Files to Folder**")
        file_path_input = st.text_input("Enter full path to file:", 
                                       placeholder="C:/Users/Scarl/Downloads/document.pdf",
                                       key="file_path_input")
        add_btn = st.button("📂 Add File", use_container_width=True)
        
        if add_btn and file_path_input:
            if os.path.isfile(file_path_input):
                try:
                    dest_path = os.path.join(folder_path, os.path.basename(file_path_input))
                    shutil.copy(file_path_input, dest_path)
                    st.success(f"✅ Added: {os.path.basename(file_path_input)}")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Failed to copy file: {e}")
            else:
                st.warning("⚠️ Invalid file path.")
    
    # File Selection
    with st.expander("🔍 Filter & Select Files", expanded=True):
        filter_keyword = st.text_input("Filter by filename:", 
                                      placeholder="Enter keyword",
                                      key="filter_keyword")
        
        file_exts = (".pdf", ".docx", ".txt", ".csv", ".xlsx")
        all_files = sorted([
            f for f in os.listdir(folder_path)
            if f.lower().endswith(file_exts) and os.path.isfile(os.path.join(folder_path, f))
        ])
        
        filtered_files = [f for f in all_files if filter_keyword.lower() in f.lower()] if filter_keyword else all_files
        
        st.markdown(f"**Available Files** ({len(filtered_files)} documents)")
        selected_files = st.multiselect("Select files to analyze:", 
                                       filtered_files, 
                                       default=st.session_state.selected_files,
                                       key="file_selector")
        
        st.session_state.selected_files = selected_files
    
    # Mode Selection
    with st.expander("⚙️ Mode Selection", expanded=True):
        mode = st.radio("Select mode:", 
                       ["📤 Single File Mode", "📂 Folder Mode"],
                       key="mode_selector")
        
        # Single file uploader
        if mode == "📤 Single File Mode":
            uploaded_file = st.file_uploader("Upload a supported file:", 
                                            type=["pdf", "docx", "txt", "csv", "xlsx"],
                                            key="file_uploader")
            
            if uploaded_file:
                # Save to temp file
                file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                # Extract content
                content = extract_text_from_file(file_path)
                st.session_state.file_preview = content
                st.session_state.uploaded_file = uploaded_file
            else:
                st.session_state.uploaded_file = None
    
    # Settings
    with st.expander("⚙️ Settings", expanded=False):
        st.info("Using Ollama with llama3 model")
        if st.button("🧹 Clear Chat History", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.last_answer = ""
            st.success("Chat history cleared!")

# === Main Content - Chat Interface ===
# Title with minimal spacing
st.markdown('<div class="title-container">', unsafe_allow_html=True)
st.title("AyF 0.1 Beta")
st.caption("Ask your Folder")
st.markdown('</div>', unsafe_allow_html=True)

# === Main Content Layout ===
st.markdown('<div class="main-container">', unsafe_allow_html=True)

# === Input Box directly below the title with minimal spacing ===
st.markdown('<div class="input-container">', unsafe_allow_html=True)

question = st.text_input("💬 Ask a question about your documents:",
                         value=st.session_state.question,
                         key="question_input")

# Create columns for buttons
st.markdown('<div class="button-columns">', unsafe_allow_html=True)
col1, col2 = st.columns(2)

with col1:
    if st.button("🤖 Ask Question", use_container_width=True) and question.strip():
        with st.spinner("🧠 Analyzing documents..."):
            docs = []
            
            if mode == "📤 Single File Mode" and st.session_state.uploaded_file:
                # Process single file
                doc = Document(page_content=st.session_state.file_preview)
                if doc:
                    docs.append(doc)
            elif mode == "📂 Folder Mode" and st.session_state.selected_files:
                # Process folder files
                for fname in st.session_state.selected_files:
                    path = os.path.join(st.session_state.folder_path, fname)
                    doc = load_doc(path)
                    if doc:
                        docs.append(doc)
            
            if docs:
                try:
                    llm = Ollama(model="llama3")
                    chain = load_qa_chain(llm, chain_type="stuff")
                    answer = chain.run(input_documents=docs, question=question)

                    st.session_state.chat_history.append(("Slim", question))
                    st.session_state.chat_history.append(("Django", answer))
                    st.session_state.last_answer = answer
                    st.session_state.play_audio = False
                    st.session_state.question = question
                    
                    # Rerun to update UI
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error: {e}")
            else:
                st.warning("⚠️ No valid documents to analyze")

with col2:
    if st.session_state.last_answer:
        if st.button("🔊 Read Answer Aloud", use_container_width=True):
            st.session_state.play_audio = True

st.markdown('</div>', unsafe_allow_html=True)  # Close button-columns
st.markdown('</div>', unsafe_allow_html=True)  # Close input-container

# === Chat History Below Input - No Space ===
st.markdown('<div class="chat-history-container">', unsafe_allow_html=True)
st.markdown('<div class="chat-history">', unsafe_allow_html=True)

if st.session_state.chat_history:
    # Display messages in reverse order (newest at top)
    for role, msg in reversed(st.session_state.chat_history):
        if role == "Slim":
            st.markdown(f"""
            <div class="message user-message">
                <div class="message-header">
                    <span>👤 Slim</span>
                </div>
                {msg}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="message assistant-message">
                <div class="message-header">
                    <span>🤖 Django</span>
                </div>
                {msg}
            </div>
            """, unsafe_allow_html=True)
else:
    st.info("Your conversation history will appear here")

st.markdown('</div>', unsafe_allow_html=True)  # Close chat-history
st.markdown('</div>', unsafe_allow_html=True)  # Close chat-history-container
st.markdown('</div>', unsafe_allow_html=True)  # Close main-container

# Audio player
if st.session_state.play_audio and st.session_state.last_answer:
    with st.container():
        st.subheader("🔊 Audio Response")
        try:
            lang_code = detect(st.session_state.last_answer)
        except:
            lang_code = "en"
        play_gtts(st.session_state.last_answer, lang_code=lang_code)
        st.session_state.play_audio = False

# File preview section
if (mode == "📤 Single File Mode" and st.session_state.uploaded_file) or \
   (mode == "📂 Folder Mode" and st.session_state.selected_files):
    st.divider()
    st.subheader("Document Preview")
    
    if mode == "📤 Single File Mode" and st.session_state.uploaded_file:
        uploaded_file = st.session_state.uploaded_file
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        file_icon = get_file_icon(file_ext)
        
        st.markdown(f"**Document:** `{file_icon} {uploaded_file.name}`")
        
        # Show preview
        st.text_area("Content Preview", 
                    value=st.session_state.file_preview[:5000], 
                    height=300,
                    disabled=True)
    
    elif mode == "📂 Folder Mode" and st.session_state.selected_files:
        # Show selected files
        st.markdown("**Selected Documents:**")
        for file in st.session_state.selected_files:
            file_ext = os.path.splitext(file)[1].lower()
            file_icon = get_file_icon(file_ext)
            st.markdown(f"- {file_icon} {file}")
